#!/usr/bin/env python3
"""Regenerate Orbit PRD Word doc from markdown."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

MD_PATH = "/Users/manudevmalik/Orbit-Student-Calendar-PRD.md"
DOCX_PATH = "/Users/manudevmalik/Orbit-Student-Calendar-PRD.docx"


def add_formatted_paragraph(doc, text, style=None, bold=False, italic=False, size=11):
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    return p


def parse_table_lines(lines):
    rows = []
    for line in lines:
        if re.match(r"^\|[\s\-:|]+\|$", line):
            continue
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        if cells:
            rows.append(cells)
    return rows


def add_table(doc, rows):
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, cell in enumerate(row):
            if j < len(table.rows[i].cells):
                table.rows[i].cells[j].text = cell
                for paragraph in table.rows[i].cells[j].paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
                        run.font.name = "Calibri"
    doc.add_paragraph()


def convert_md_to_docx(md_path, docx_path):
    with open(md_path, "r", encoding="utf-8") as f:
        content = f.read()

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    lines = content.split("\n")
    i = 0
    in_code_block = False
    code_lines = []
    table_lines = []

    while i < len(lines):
        line = lines[i]

        if line.strip().startswith("```"):
            if in_code_block:
                code_text = "\n".join(code_lines)
                p = doc.add_paragraph()
                run = p.add_run(code_text)
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                p.paragraph_format.left_indent = Inches(0.25)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        if line.strip().startswith("|"):
            table_lines.append(line)
            i += 1
            if i >= len(lines) or not lines[i].strip().startswith("|"):
                rows = parse_table_lines(table_lines)
                add_table(doc, rows)
                table_lines = []
            continue

        if line.startswith("# "):
            p = doc.add_heading(line[2:].strip(), level=0)
            for run in p.runs:
                run.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)
            i += 1
            continue

        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
            i += 1
            continue

        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
            i += 1
            continue

        if line.startswith("#### "):
            doc.add_heading(line[5:].strip(), level=3)
            i += 1
            continue

        if line.strip() == "---":
            doc.add_paragraph()
            i += 1
            continue

        if line.strip().startswith(">"):
            text = line.strip().lstrip(">").strip()
            p = add_formatted_paragraph(doc, text, italic=True, size=11)
            p.paragraph_format.left_indent = Inches(0.3)
            i += 1
            continue

        if line.strip().startswith("- "):
            text = line.strip()[2:]
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
            p = doc.add_paragraph(text, style="List Bullet")
            for run in p.runs:
                run.font.size = Pt(10)
                run.font.name = "Calibri"
            i += 1
            continue

        if re.match(r"^\d+\.\s", line.strip()):
            text = re.sub(r"^\d+\.\s", "", line.strip())
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
            p = doc.add_paragraph(text, style="List Number")
            for run in p.runs:
                run.font.size = Pt(10)
                run.font.name = "Calibri"
            i += 1
            continue

        if line.strip() == "":
            i += 1
            continue

        text = line.strip()
        text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
        text = re.sub(r"\*(.+?)\*", r"\1", text)
        text = text.replace("`", "")

        if text.startswith("**") and text.endswith("**"):
            add_formatted_paragraph(doc, text.strip("*"), bold=True)
        else:
            p = doc.add_paragraph()
            parts = re.split(r"(\*\*.+?\*\*)", text)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                elif part:
                    run = p.add_run(part)
                run.font.size = Pt(10)
                run.font.name = "Calibri"

        i += 1

    doc.save(docx_path)
    print(f"Saved: {docx_path}")


if __name__ == "__main__":
    convert_md_to_docx(MD_PATH, DOCX_PATH)
